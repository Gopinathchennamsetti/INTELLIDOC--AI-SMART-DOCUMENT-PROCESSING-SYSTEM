import os
import docx
import logging
from openai_utils import identify_duplicates, deduplicate_text, summarize_text


def read_docx(file_path: str) -> str:
    """
    Read and return text from a .docx file.
    """
    try:
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        text = "\n".join(full_text)
        logging.info(f"read_docx: {file_path} -> {len(text)} chars")
        return text
    except Exception as e:
        logging.error(f"Error reading document {file_path}: {e}")
        raise RuntimeError(f"Error reading document: {str(e)}")


def save_docx(file_path: str, text: str):
    """
    Save text to a .docx file.
    """
    try:
        doc = docx.Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)
        doc.save(file_path)
        logging.info(f"save_docx: {file_path} <- {len(text)} chars")
    except Exception as e:
        logging.error(f"Error saving document {file_path}: {e}")
        raise RuntimeError(f"Error saving document: {str(e)}")


def load_documents_from_filepaths(filepaths):
    """
    Load and return documents and filenames from given file paths.
    """
    documents = []
    filenames = []
    for filepath in filepaths:
        if filepath.endswith('.docx'):
            try:
                doc = docx.Document(filepath)
                text = "\n".join([para.text for para in doc.paragraphs])
                documents.append(text)
                filenames.append(os.path.basename(filepath))
                logging.info(f"Document loaded successfully: {filepath}")
            except Exception as e:
                error_message = f"Error loading document {os.path.basename(filepath)}: {str(e)}"
                logging.error(error_message)
                documents.append(error_message)
    return documents, filenames


def process_single_document(file_path, duplicate_folder, consolidated_folder, summary_folder):
    """
    Process a document: identify duplicates, deduplicate, consolidate, and summarize.
    Uses Gemini through openai_utils.
    """
    try:
        logging.info(f'Processing document: {file_path}')
        document_text = read_docx(file_path)
        logging.info(f'Original document length: {len(document_text)} chars')

        # Identify and save duplicates
        duplicates_text = identify_duplicates(document_text).strip()
        logging.info(f'Duplicates text length: {len(duplicates_text)} chars')
        if duplicates_text:
            duplicates_path = os.path.join(duplicate_folder, os.path.basename(file_path))
            save_docx(duplicates_path, duplicates_text)

        # Deduplicate and save consolidated text
        deduped_text = deduplicate_text(document_text)
        logging.info(f'Deduplicated text length: {len(deduped_text)} chars')

        consolidated_text = f"Consolidated Content:\n\n{deduped_text}"
        consolidated_path = os.path.join(consolidated_folder, os.path.basename(file_path))
        save_docx(consolidated_path, consolidated_text)
        logging.info(f"Saved consolidated file: {consolidated_path}")

        # Summarize and save summary
        summary_text = summarize_text(deduped_text)
        logging.info(f'Summary text length: {len(summary_text)} chars')
        summary_path = os.path.join(summary_folder, os.path.basename(file_path))
        save_docx(summary_path, summary_text)
        logging.info(f"Document processing completed: {file_path}")

    except Exception as e:
        logging.error(f"Error processing document {file_path}: {e}")
        raise RuntimeError(f"Error processing document: {str(e)}")