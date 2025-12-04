import os
import re
import logging
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import AffinityPropagation
from sklearn.metrics import silhouette_score
from flask import render_template, redirect, request
from werkzeug.utils import secure_filename
from config import UPLOAD_FOLDER, CONSOLIDATED_FOLDER, CLUSTERED_FOLDER
from file_utils import load_documents_from_filepaths
from text_processing import preprocess_text
from docx import Document


def cluster_documents(request):
    """
    Handles document clustering: processes files, clusters documents, and renders results.
    """
    if request.method == 'POST':
        if 'files' not in request.files:
            return redirect(request.url)

        files = request.files.getlist('files')
        filepaths = []

        # Save uploaded files
        try:
            for file in files:
                filepath = os.path.join(UPLOAD_FOLDER, secure_filename(file.filename))
                file.save(filepath)
                filepaths.append(filepath)
            logging.info(f'Files saved: {filepaths}')
        except Exception as e:
            logging.error(f'Error saving files: {e}')
            return render_template('error.html', error_message="An error occurred while saving files.")

        try:
            # Load and preprocess
            documents, filenames = load_documents_from_filepaths(filepaths)
            df = pd.DataFrame({'filename': filenames, 'text': documents})
            df['cleaned_text'] = df['text'].apply(preprocess_text)

            # TF-IDF
            vectorizer = TfidfVectorizer()
            X = vectorizer.fit_transform(df['cleaned_text'])

            # Clustering
            affinity_propagation = AffinityPropagation()
            clusters = affinity_propagation.fit_predict(X)

            df['cluster'] = clusters

            # Silhouette
            silhouette_avg = None
            if len(set(clusters)) > 1:
                silhouette_avg = silhouette_score(X, clusters, metric='cosine')
                logging.info(f'Silhouette Score: {silhouette_avg:.3f}')
            else:
                logging.info('Silhouette Score not computed (single cluster).')

            # Simple cluster names
            cluster_names = []
            for cluster in set(clusters):
                cluster_name = f'Company {cluster + 1}'
                cluster_names.append(cluster_name)

            df['cluster_name'] = df['cluster'].apply(lambda x: cluster_names[x])

            # Save CSV
            output_csv_path = os.path.join(CONSOLIDATED_FOLDER, 'clustered_documents_affinity_propagation.csv')
            df.to_csv(output_csv_path, index=False)
            logging.info(f'Results saved to {output_csv_path}')

            # Create consolidated docs for each cluster
            for i, cluster_name in enumerate(cluster_names):
                doc = Document()
                doc.add_heading(f'Cluster {i + 1} ({cluster_name})', level=1)
                cluster_docs = df[df['cluster'] == i]['text'].tolist()
                combined_text = "\n\n".join(cluster_docs)
                combined_text_no_citations = re.sub(r'\[[^\]]*\]', '', combined_text)
                doc.add_paragraph(combined_text_no_citations)
                output_docx_path = os.path.join(CLUSTERED_FOLDER, f'{cluster_name}_Consolidated.docx')
                doc.save(output_docx_path)
                logging.info(f'Consolidated doc saved: {output_docx_path}')

            results = []
            for cluster in set(clusters):
                cluster_data = {
                    'id': cluster,
                    'name': cluster_names[cluster],
                    'documents': df[df['cluster'] == cluster][['filename', 'text']].to_dict(orient='records'),
                    'consolidated_filename': f'{cluster_names[cluster]}_Consolidated.docx'
                }
                results.append(cluster_data)

            return render_template('index.html', clusters=results, silhouette_score=silhouette_avg)

        except Exception as e:
            logging.error(f'Error during clustering or processing: {e}')
            return render_template('error.html', error_message="An error occurred during clustering or file processing.")

    return render_template('index.html', clusters=None)