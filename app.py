# app.py
# Importing all the necessary libraries
import os
import re
import logging
import numpy as np
import pandas as pd
import spacy
import nltk
from collections import Counter
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, abort
from werkzeug.utils import secure_filename
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.decomposition import PCA
from sklearn.cluster import AffinityPropagation
from sklearn.metrics import silhouette_score
import plotly.express as px
import plotly.io as pio

# Import Gemini utilities
from openai_utils import identify_duplicates, deduplicate_text, summarize_text

# ---------- Setup ----------
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s %(levelname)s: %(message)s',
    handlers=[logging.FileHandler("app.log"), logging.StreamHandler()]
)

# NLP setup
nlp = spacy.load('en_core_web_sm')
nltk.download('stopwords')
nltk.download('punkt')

# Flask app setup
app = Flask(__name__)
app.secret_key = '12345678'

# Folders
UPLOAD_FOLDER = 'uploads'
CLUSTERED_FOLDER = 'clustered_output'
CONSOLIDATED_FOLDER = 'consolidated_output'
DUPLICATE_FOLDER = 'duplicate_output'
SUMMARY_FOLDER = 'summary_output'

for folder in [UPLOAD_FOLDER, CLUSTERED_FOLDER, CONSOLIDATED_FOLDER, DUPLICATE_FOLDER, SUMMARY_FOLDER]:
    os.makedirs(folder, exist_ok=True)

# ---------- Helper Functions ----------
def load_documents_from_filepaths(filepaths):
    documents, filenames = [], []
    for filepath in filepaths:
        try:
            if filepath.endswith('.docx'):
                doc = Document(filepath)
                text = "\n".join([para.text for para in doc.paragraphs])
                documents.append(text)
                filenames.append(os.path.basename(filepath))
        except Exception as e:
            logging.error(f"Error processing file {filepath}: {e}")
    return documents, filenames

def preprocess_text(text):
    try:
        text = re.sub(r'\[[^\]]*\]', '', text.lower())
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s]', '', text)
        doc = nlp(text)
        return ' '.join([token.lemma_ for token in doc if not token.is_stop and not token.is_punct])
    except Exception as e:
        logging.error(f"Error preprocessing text: {e}")
        return ""

def identify_company_name(texts):
    all_entities = []
    try:
        for text in texts:
            doc = nlp(text)
            entities = [ent.text for ent in doc.ents if ent.label_ == 'ORG']
            all_entities.extend(entities)
        if all_entities:
            most_common = Counter(all_entities).most_common(1)
            return most_common[0][0]
    except Exception as e:
        logging.error(f"Error identifying company name: {e}")
    return 'Unknown'

def read_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logging.error(f"Error reading file {file_path}: {e}")
        raise

def save_docx(file_path, text):
    try:
        doc = Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)
        doc.save(file_path)
    except Exception as e:
        logging.error(f"Error saving file {file_path}: {e}")
        raise

def process_single_document(file_path, duplicate_folder, consolidated_folder, summary_folder):
    logging.info(f"Processing document: {file_path}")
    document_text = read_docx(file_path)

    # Identify duplicates
    duplicates_text = identify_duplicates(document_text)
    if duplicates_text:
        duplicates_path = os.path.join(duplicate_folder, os.path.basename(file_path))
        save_docx(duplicates_path, duplicates_text)
        logging.info(f"Saved duplicates at: {duplicates_path}")

    # Deduplicate
    deduped_text = deduplicate_text(document_text)
    consolidated_text = f"Consolidated Content:\n\n{deduped_text}"
    consolidated_path = os.path.join(consolidated_folder, os.path.basename(file_path))
    save_docx(consolidated_path, consolidated_text)
    logging.info(f"Saved consolidated at: {consolidated_path}")

    # Summarize
    summary_text = summarize_text(deduped_text)
    summary_path = os.path.join(summary_folder, os.path.basename(file_path))
    save_docx(summary_path, summary_text)
    logging.info(f"Saved summary at: {summary_path}")

def compute_bertscore(ref_text, hyp_text):
    from bert_score import score as bert_score
    try:
        P, R, F1 = bert_score([hyp_text], [ref_text], lang='en', verbose=True)
        return P.mean().item(), R.mean().item(), F1.mean().item()
    except Exception as e:
        logging.error(f"Error computing BERTScore: {e}")
        return 0, 0, 0

def compute_cosine_similarity(ref_text, hyp_text):
    try:
        vectorizer = TfidfVectorizer().fit_transform([ref_text, hyp_text])
        vectors = vectorizer.toarray()
        return cosine_similarity(vectors)[0, 1]
    except Exception as e:
        logging.error(f"Error computing cosine similarity: {e}")
        return 0

def compare_folders(ref_folder, hyp_folder, output_csv):
    results = []
    try:
        ref_files = sorted(os.listdir(ref_folder))
        hyp_files = sorted(os.listdir(hyp_folder))
        for ref_file, hyp_file in zip(ref_files, hyp_files):
            if ref_file.endswith('.docx') and hyp_file.endswith('.docx'):
                ref_path = os.path.join(ref_folder, ref_file)
                hyp_path = os.path.join(hyp_folder, hyp_file)
                reference_text = read_docx(ref_path)
                hypothesis_text = read_docx(hyp_path)
                P, R, F1 = compute_bertscore(reference_text, hypothesis_text)
                cosine_sim = compute_cosine_similarity(reference_text, hypothesis_text)
                results.append({
                    'File': ref_file,
                    'BERTScore Precision': P,
                    'BERTScore Recall': R,
                    'BERTScore F1': F1,
                    'Cosine Similarity': cosine_sim
                })
        pd.DataFrame(results).to_csv(output_csv, index=False)
        logging.info(f"Results saved to: {output_csv}")
    except Exception as e:
        logging.error(f"Error comparing folders: {e}")

# ---------- Flask Routes ----------
@app.route('/', methods=['GET', 'POST'])
def index():
    return render_template('home.html', clusters=None)

@app.route('/cluster', methods=['GET', 'POST'])
def cluster():
    if request.method == 'POST':
        try:
            if 'files' not in request.files:
                flash('No files part in the request')
                return redirect(request.url)
            files = request.files.getlist('files')
            filepaths = [os.path.join(UPLOAD_FOLDER, secure_filename(f.filename)) for f in files]
            for f, path in zip(files, filepaths):
                f.save(path)

            documents, filenames = load_documents_from_filepaths(filepaths)
            df = pd.DataFrame({'filename': filenames, 'text': documents})
            df['cleaned_text'] = df['text'].apply(preprocess_text)

            X = TfidfVectorizer().fit_transform(df['cleaned_text'])
            clusters = AffinityPropagation().fit_predict(X)
            df['cluster'] = clusters
            silhouette_avg = silhouette_score(X, clusters, metric='cosine') if len(set(clusters)) > 1 else None

            cluster_names = []
            for cluster in set(clusters):
                cluster_docs = df[df['cluster'] == cluster]['text'].tolist()
                name = identify_company_name(cluster_docs)
                cluster_names.append(name if name != 'Unknown' else f'Company {cluster + 1}')
            df['cluster_name'] = df['cluster'].apply(lambda x: cluster_names[x])

            # Save CSV
            df.to_csv(os.path.join(CONSOLIDATED_FOLDER, 'clustered_documents_affinity_propagation.csv'), index=False)

            # Create consolidated docx for each cluster
            for i, cluster_name in enumerate(cluster_names):
                doc = Document()
                doc.add_heading(f'Cluster {i + 1} ({cluster_name})', level=1)
                cluster_docs = df[df['cluster'] == i]['text'].tolist()
                doc.add_paragraph("\n\n".join([re.sub(r'\[[^\]]*\]', '', t) for t in cluster_docs]))
                doc.save(os.path.join(CLUSTERED_FOLDER, f'{cluster_name}_Consolidated.docx'))

            # Plotly scatter
            X_reduced = PCA(n_components=2).fit_transform(X.toarray())
            fig = px.scatter(
                x=X_reduced[:, 0],
                y=X_reduced[:, 1],
                color=df['cluster_name'],
                hover_name=df['filename'],
                labels={'x':'PCA 1','y':'PCA 2','color':'Company'},
                title='Document Clustering Visualization',
                color_discrete_sequence=px.colors.qualitative.Plotly
            )
            plot_html = pio.to_html(fig, full_html=False)
            with open(os.path.join(CONSOLIDATED_FOLDER, 'cluster_viz.html'), 'w', encoding='utf-8') as f:
                f.write(plot_html)

            results = [{'id': cluster, 'name': cluster_names[cluster],
                        'documents': df[df['cluster'] == cluster][['filename','text']].to_dict('records')}
                       for cluster in set(clusters)]

            return render_template('index.html', clusters=results, silhouette_score=silhouette_avg,
                                   plot_url=url_for('plot_viz'))
        except Exception as e:
            logging.error(f"Error in clustering: {e}")
            return render_template('error.html', message='An error occurred during clustering.')

    return render_template('index.html', clusters=None)

@app.route('/plot_viz')
def plot_viz():
    return send_from_directory(CONSOLIDATED_FOLDER, 'cluster_viz.html')

# Consolidate
@app.route('/consolidate', methods=['GET','POST'])
def consolidate():
    if request.method == 'POST':
        try:
            selected_file = request.form.get('document')
            if selected_file:
                file_path = os.path.join(CLUSTERED_FOLDER, selected_file)
                process_single_document(file_path, DUPLICATE_FOLDER, CONSOLIDATED_FOLDER, SUMMARY_FOLDER)
                flash('Consolidation completed!')
                return redirect(request.url)
            else:
                flash('No document selected.')
        except Exception as e:
            flash(f'Error: {e}')
            app.logger.error(f'Error in consolidation: {e}')
            return redirect(request.url)
    files = [f for f in os.listdir(CLUSTERED_FOLDER) if f.endswith('.docx')]
    return render_template('consolidate.html', files=files, message=None)

# Summary
@app.route('/summary', methods=['GET','POST'])
def summary():
    if request.method == 'POST':
        selected_file = request.form.get('document')
        if selected_file:
            try:
                file_path = os.path.join(CLUSTERED_FOLDER, selected_file)
                document_text = read_docx(file_path)
                summary_text = summarize_text(document_text)
                save_docx(os.path.join(SUMMARY_FOLDER, selected_file), summary_text)
                return redirect(request.url)
            except Exception as e:
                return f"Error processing file {selected_file}: {e}", 500
    files = [f for f in os.listdir(CLUSTERED_FOLDER) if f.endswith('.docx')]
    return render_template('summary.html', files=files, message=None)

# Evaluate
@app.route('/evaluate')
def compare_folders_route():
    try:
        output_csv = os.path.join(CONSOLIDATED_FOLDER, 'comparison_results.csv')
        compare_folders(CLUSTERED_FOLDER, CONSOLIDATED_FOLDER, output_csv)
        return redirect(url_for('show_results'))
    except Exception as e:
        return render_template('error.html', error_message=f"Error: {e}"), 500

@app.route('/show_results', methods=['GET'])
def show_results():
    csv_file = os.path.join(CONSOLIDATED_FOLDER, 'comparison_results.csv')
    if not os.path.isfile(csv_file):
        return render_template('error.html', error_message="Results file not found."), 404
    df = pd.read_csv(csv_file)
    return render_template('results.html', tables=[df.to_html(classes='data')], titles=df.columns.values)

# View files
@app.route('/view_clusters')
def view_clusters():
    files = [f for f in os.listdir(CLUSTERED_FOLDER) if f.endswith('.docx')]
    return render_template('view_clusters.html', files=files)

@app.route('/view_consolidation')
def view_consolidation():
    files = [f for f in os.listdir(CONSOLIDATED_FOLDER) if f.endswith('.docx')]
    return render_template('view_consolidation.html', files=files)

@app.route('/view_summary')
def view_summary():
    files = [f for f in os.listdir(SUMMARY_FOLDER) if f.endswith('.docx')]
    return render_template('view_summary.html', files=files)

@app.route('/cluster_file/<filename>')
def cluster_file(filename):
    file_path = os.path.join(CLUSTERED_FOLDER, filename)
    if not os.path.isfile(file_path):
        abort(404)
    text = read_docx(file_path)
    return render_template('view_file.html', content=text)

@app.route('/consolidated_file/<filename>')
def consolidated_file(filename):
    file_path = os.path.join(CONSOLIDATED_FOLDER, filename)
    if not os.path.isfile(file_path):
        abort(404)
    text = read_docx(file_path)
    return render_template('view_file.html', content=text)

@app.route('/summary_file/<filename>')
def summary_file(filename):
    file_path = os.path.join(SUMMARY_FOLDER, filename)
    if not os.path.isfile(file_path):
        abort(404)
    text = read_docx(file_path)
    return render_template('view_file.html', content=text)

@app.route('/download_clustered/<filename>')
def download_clustered(filename):
    return send_from_directory(CLUSTERED_FOLDER, filename, as_attachment=True)

@app.route('/download_consolidated/<filename>')
def download_consolidated(filename):
    return send_from_directory(CONSOLIDATED_FOLDER, filename, as_attachment=True)

@app.route('/download_summary/<filename>')
def download_summary(filename):
    return send_from_directory(SUMMARY_FOLDER, filename, as_attachment=True)

# ---------- Main ----------
if __name__ == '__main__':
    app.run(debug=True)
